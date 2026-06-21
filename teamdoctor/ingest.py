"""Pull plain text out of an uploaded description file.

Supports txt/md (decode), PDF (pypdf), and Word .docx (python-docx). Legacy
binary .doc has no clean pure-Python reader, so it gets a best-effort fallback —
if that yields too little, the caller tells the user to use PDF or .docx.

Keeps the app dependency-light and avoids AGPL PDF libraries (pypdf is BSD).
"""

from __future__ import annotations

import io
import re

MIN_USABLE = 20  # below this many chars, treat extraction as failed


def extract_text(uploaded) -> str:
    """uploaded is a Streamlit UploadedFile. Returns extracted plain text."""
    name = (getattr(uploaded, "name", "") or "").lower()
    data = uploaded.read()
    if name.endswith((".txt", ".md")):
        return data.decode("utf-8", errors="ignore")
    if name.endswith(".pdf"):
        return _pdf(data)
    if name.endswith(".docx"):
        return _docx(data)
    if name.endswith(".doc"):
        return _doc_fallback(data)
    return data.decode("utf-8", errors="ignore")


def _pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages).strip()


def _docx(data: bytes) -> str:
    import docx
    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return "\n".join(t for t in parts if t.strip()).strip()


def _doc_fallback(data: bytes) -> str:
    """Legacy binary .doc — best-effort: keep runs of printable text."""
    raw = data.decode("latin-1", errors="ignore")
    chunks = re.findall(r"[\x20-\x7e][\x20-\x7e\r\n\t]{3,}", raw)
    text = " ".join(c.strip() for c in chunks)
    return re.sub(r"\s{2,}", " ", text).strip()
